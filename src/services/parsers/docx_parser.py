"""DOCX document parser."""
from pathlib import Path
from typing import Dict, Any
import docx
from src.services.parsers.base import BaseParser
from src.utils.logger import logger
from src.utils.errors import DocumentParseError


class DOCXParser(BaseParser):
    """Parse Microsoft Word documents."""
    
    async def parse(self, file_path: Path) -> Dict[str, Any]:
        """Parse DOCX document."""
        try:
            logger.info(f"Parsing DOCX: {file_path.name}")
            
            doc = docx.Document(file_path)
            
            # Extract text content
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            content = '\n\n'.join(paragraphs)
            
            # Extract metadata
            core_props = doc.core_properties
            metadata = {
                'title': core_props.title or '',
                'author': core_props.author or '',
                'created': str(core_props.created) if core_props.created else '',
                'modified': str(core_props.modified) if core_props.modified else '',
                'paragraphs': len(doc.paragraphs),
                'sections': len(doc.sections),
            }
            
            return {
                'content': content,
                'metadata': metadata,
                'format': 'docx'
            }
            
        except Exception as e:
            logger.error(f"Failed to parse DOCX {file_path}: {e}")
            raise DocumentParseError(f"DOCX parse error: {e}")
    
    def supports_format(self, file_extension: str) -> bool:
        """Check if format is supported."""
        return file_extension.lower() in ['.docx', '.doc']
